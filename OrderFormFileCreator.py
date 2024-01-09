import docx
from docx import Document
import openpyxl
from openpyxl import load_workbook

'''
This program is a script that produces a mochi order form as a word document 
using a pre-created template

author: J Kataoka
'''


# Constants for the Mochi Prices and the date
komochiPrice = 8.00
okagamiPrice = 8.50
ankoPrice = 12.00
date = "12/30/2023"

# Template File Name and Excel Datasheet Name
orderFormTemplate = 'Order Form Template.docx'
orderFormData = 'Example Orders.xlsx'

# Small function for generic text formatting
def defaultFormats(inputRun):
    inputRun.bold = True
    inputRun.font.name = 'Arial Black'
    inputRun.font.size = docx.shared.Pt(12)

def orderMaker(name,address,phone,komochi,okagami,anko):
    
    '''
    Produces a word document of a Mochi order form/reciept

        :param name: name of the customer
        :param phone: phone number of the customer
        :param address: email or road address of customer
        :param komochi: the number of komochi ordered (Integer)
        :param okagami: the number of okagami mochi ordered (Integer)
        :param anko: the number of anko mochi ordered (Integer)

    '''

    # Checks for certain cases
    if name is None:
        raise Exception("Error, Name must be given/not NULL")

    if komochi is None:
        komochi = 0
    
    if okagami is None:
        okagami = 0

    if anko is None:
        anko = 0

    # Opens the pre-created template word document
    doc = Document(orderFormTemplate)

    # 14 columns, 19 rows.
    table = doc.tables[0]

    cell = table.rows[3].cells[1] # Cell for name 
    p = cell.paragraphs[0] # Accesses the existing cell paragraph
    run = p.add_run(name) # Add person's name
    defaultFormats(run) # Apply default formatting

    cell = table.rows[3].cells[12] 
    p = cell.paragraphs[0] 
    run = p.add_run(phone) 
    defaultFormats(run) 

    cell = table.rows[5].cells[1] 
    p = cell.paragraphs[0] 
    run = p.add_run(address) 
    defaultFormats(run) 

    cell = table.rows[7].cells[0] 
    p = cell.paragraphs[0] 
    run = p.add_run("(   " + str(komochi) + "   )") 
    defaultFormats(run) 

    cell = table.rows[9].cells[0]  
    p = cell.paragraphs[0] 
    run = p.add_run("(   " + str(okagami) + "   )") 
    defaultFormats(run) 

    cell = table.rows[11].cells[0] 
    p = cell.paragraphs[0] 
    run = p.add_run("(   " + str(anko) + "   )") 
    defaultFormats(run) 

    cell = table.rows[7].cells[10] 
    p = cell.paragraphs[0] 
    run = p.add_run("$" + str(komochi*komochiPrice)) 
    defaultFormats(run) 

    cell = table.rows[9].cells[10] 
    p = cell.paragraphs[0] 
    run = p.add_run("$" + str(okagami*okagamiPrice)) 
    defaultFormats(run)

    cell = table.rows[11].cells[10] 
    p = cell.paragraphs[0] 
    run = p.add_run("$" + str(anko*ankoPrice)) 
    defaultFormats(run) 

    total = komochi*komochiPrice + okagami*okagamiPrice + anko*ankoPrice
    cell = table.rows[14].cells[10] 
    p = cell.paragraphs[0] 
    run = p.add_run("$" + str(total)) 
    defaultFormats(run) 

    cell = table.rows[16].cells[10] 
    p = cell.paragraphs[0] 
    run = p.add_run(date) 
    defaultFormats(run) 

    doc.save(name + ".docx")


# MAIN METHOD

wb = load_workbook(filename = orderFormData)
ws = wb.active

data = []

for row in range(4,ws.max_row):
    if(ws.cell(row,1).value is None):
        break

    for i in range(1,7):
        data.append(ws.cell(row,i).value)

    orderMaker(data[0],data[1],data[2],data[3],data[4],data[5])
    data = []